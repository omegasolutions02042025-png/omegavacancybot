# -*- coding: utf-8 -*-
"""
White Label Resume Builder:
1) build_white_label_prompt(...) — промпт для GPT/Gemini (строгие правила, Projects — обязательно).
2) generate_resume_payload_gemini(...) — запрос к Gemini, возврат JSON payload {config, content}.
3) render_resume_docx(payload) — рендер красивого .docx по JSON (Times New Roman, синие заголовки).
4) create_white_label_resume(...) — полный конвейер: кандидатский текст -> JSON -> DOCX.
5) parse_json_loose(...) — «живучий» парсер JSON из свободного текста модели.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK
import json, re

def _extract_text_from_gemini_response(resp) -> str:
    try:
        if getattr(resp, "text", None):
            return (resp.text or "").strip()
    except Exception:
        pass
    out = []
    try:
        for cand in (getattr(resp, "candidates", None) or []):
            content = getattr(cand, "content", None)
            parts = getattr(content, "parts", None) if content else None
            if not parts:
                continue
            for p in parts:
                t = getattr(p, "text", None)
                if t: out.append(t)
    except Exception:
        pass
    return "\n".join(out).strip()

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = (hex_color or "#000000").lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _set_core_styles(doc: Document, font_family: str, font_size_main: int):
    style = doc.styles['Normal']
    style.font.name = font_family
    style.font.size = Pt(font_size_main)
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

def _add_section_title(doc: Document, title: str, color_hex: str, font_size_headings: int):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(font_size_headings)
    r.font.color.rgb = _hex_to_rgb(color_hex)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

def _add_text(doc: Document, text: str, bold: bool = False):
    if not text:
        return
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].bold = bool(bold)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(2)

def _render_skills(doc: Document, skills):
    if not skills:
        return
    if isinstance(skills, str):
        _add_text(doc, skills)
    elif isinstance(skills, list):
        _add_text(doc, ", ".join(map(str, skills)))
    elif isinstance(skills, dict):
        for k, v in skills.items():
            line = f"{k}: {', '.join(map(str, v)) if isinstance(v, list) else v}"
            _add_text(doc, line)

def _render_experience(doc: Document, exp_list: list):
    if not exp_list:
        return
    for it in exp_list:
        header = " — ".join([x for x in [it.get("company"), it.get("position")] if x])
        if header: _add_text(doc, header, bold=True)
        if it.get("period"): _add_text(doc, it["period"], bold=True)  # Сделать период полужирным
        for key in ("responsibilities", "achievements"):
            for ln in (it.get(key) or []):
                _add_text(doc, ln)
        techs = it.get("technologies") or []
        if techs:
            _add_text(doc, f"Технологии: {', '.join(map(str, techs))}")

def _render_education(doc: Document, education):
    if not education:
        return
    if isinstance(education, str):
        _add_text(doc, education); return
    for ed in education:
        line = " — ".join(filter(None, [ed.get("institution"), ed.get("degree")]))
        if line: _add_text(doc, line, bold=True)
        if ed.get("years"): _add_text(doc, ed["years"])
        if ed.get("details"): _add_text(doc, ed["details"])

def _render_projects(doc: Document, projects: list):
    if not projects:
        return
    for pr in projects:
        if pr.get("title"): _add_text(doc, pr["title"], bold=True)
        if pr.get("role"): _add_text(doc, f"Роль: {pr['role']}")
        if pr.get("period"): _add_text(doc, f"Период: {pr['period']}")
        if pr.get("description"): _add_text(doc, pr["description"])
        techs = pr.get("technologies") or []
        if techs: _add_text(doc, f"Технологии: {', '.join(map(str, techs))}")
        if pr.get("results"): _add_text(doc, f"Результаты: {pr['results']}")

def _post_fix_bold_skills(doc: Document):
    SECTION_HEADERS = {
        "РЕЗЮМЕ", "КРАТКОЕ ОПИСАНИЕ ПРОФИЛЯ",
        "КЛЮЧЕВЫЕ НАВЫКИ", "ОПЫТ РАБОТЫ", "ОБРАЗОВАНИЕ",
        "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ", "ПРОЕКТЫ",
    }
    start_idx, end_idx = None, None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "КЛЮЧЕВЫЕ НАВЫКИ":
            start_idx = i + 1
            break
    if start_idx is None:
        return
    for j in range(start_idx, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip().upper() in SECTION_HEADERS:
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    for k in range(start_idx, end_idx):
        p = doc.paragraphs[k]
        txt = p.text
        if not txt or ":" not in txt:
            continue
        idx = txt.find(":")
        head = txt[: idx + 1].strip()
        tail = txt[idx + 1 :].lstrip()
        for r in p.runs: r.text = ""
        run_head = p.add_run(head + (" " if tail else "")); run_head.bold = True
        if tail:
            run_tail = p.add_run(tail); run_tail.bold = False

def _post_fix_inline_dicts(doc: Document):
    pattern = re.compile(r"^\s*\{\s*'title'\s*:\s*'([^']*)'\s*,\s*'items'\s*:\s*\[([^\]]*)\]\s*\}\s*$")
    for p in doc.paragraphs:
        m = pattern.match(p.text.strip())
        if not m:
            continue
        title = m.group(1).strip()
        items_raw = m.group(2).strip()
        parts = [x.strip().strip("'").strip('"') for x in items_raw.split(",") if x.strip()]
        for r in p.runs: r.text = ""
        if title:
            rt = p.add_run(title); rt.bold = True
        for it in parts:
            br = p.add_run(); br.add_break(WD_BREAK.LINE)
            p.add_run(it)

# --- Новое: форматирование даты/времени по-русски и вставка раздела «Примечание» ---

_MONTHS_RU_GEN = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
}

def _format_dt_ru(dt: datetime) -> str:
    """16 августа 2025 в 11:31"""
    return f"{dt.day} {_MONTHS_RU_GEN[dt.month]} {dt.year} в {dt.strftime('%H:%M')}"

def _render_primichanie(doc: Document, color_hex: str, font_size_headings: int, utochnenie):
    """Добавляет раздел «Примечание», если utochnenie не пусто.
       Первая строка: «Резюме обновлено: <дата> (данные из оригинала).»
       Далее: по одному «Добавлено: <элемент>» для каждого пункта из utochnenie.
    """
    # Нормализуем utochnenie к списку строк
    items = []
    if isinstance(utochnenie, str):
        s = utochnenie.strip()
        if s:
            # Разрешим разделение по точке с запятой/переводам строк/запятым — но без выкидывания смысла
            candidates = re.split(r"[;\n]+", s)
            for c in candidates:
                cc = c.strip().strip("-–•").strip()
                if cc:
                    items.append(cc)
    elif isinstance(utochnenie, (list, tuple, set)):
        for x in utochnenie:
            xx = (str(x) if not isinstance(x, str) else x).strip().strip("-–•").strip()
            if xx:
                items.append(xx)

    if not items:
        return

    _add_section_title(doc, "Примечание", color_hex, font_size_headings)
    _add_text(doc, f"Резюме обновлено: {_format_dt_ru(datetime.now())} (данные из оригинала).")
    for it in items:
        _add_text(doc, f"Добавлено: {it}")

def render_resume_docx(payload: dict, vacancy_text: str = "", utochnenie=None) -> str:
    cfg = payload.get("config", {})
    cnt = payload.get("content", {})
    doc = Document()
    _set_core_styles(doc, cfg.get("font_family", "Times New Roman"), int(cfg.get("font_size_main", 12)))
    color = cfg.get("color_headings", "#1F4E79")
    hsize = int(cfg.get("font_size_headings", 14))
    sections = cfg.get("sections", [
        "ФИО","РЕЗЮМЕ","Краткое описание профиля",
        "Ключевые навыки","Опыт работы","Образование","Дополнительная информация","Проекты"
    ])
    fio = cnt.get("fio") or {}
    # Убираем отдельное отображение ФИО в начале - оно будет в секции РЕЗЮМЕ
    
    for sec in sections:
        su = sec.upper()
        if su == "ФИО":
            continue
        _add_section_title(doc, sec, color, hsize)
        if su == "РЕЗЮМЕ":
            if fio.get("full_name"): _add_text(doc, f"ФИО: {fio['full_name']}", bold=True)
            if cnt.get("position_grade"): _add_text(doc, f"ДОЛЖНОСТЬ: {cnt['position_grade']}", bold=True)
            if cnt.get("grade"): _add_text(doc, f"Грейд: {cnt['grade']}", bold=True)
            if fio.get("location"): _add_text(doc, f"Локация: {fio['location']}", bold=True)
            if fio.get("citizenship"): _add_text(doc, f"Гражданство: {fio['citizenship']}", bold=True)
            if fio.get("birth_date"): _add_text(doc, f"Дата рождения: {fio['birth_date']}", bold=True)
        elif su == "КРАТКОЕ ОПИСАНИЕ ПРОФИЛЯ":
            if cnt.get("summary"): _add_text(doc, cnt["summary"])
        elif su == "КЛЮЧЕВЫЕ НАВЫКИ":
            _render_skills(doc, cnt.get("skills"))
        elif su == "ОПЫТ РАБОТЫ":
            _render_experience(doc, cnt.get("experience"))
        elif su == "ОБРАЗОВАНИЕ":
            _render_education(doc, cnt.get("education"))
        elif su == "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ":
            extra = cnt.get("extra")
            if isinstance(extra, list):
                for ln in extra: _add_text(doc, str(ln))
            elif extra: _add_text(doc, str(extra))
        elif su == "ПРОЕКТЫ":
            _render_projects(doc, cnt.get("projects"))
    _post_fix_bold_skills(doc)
    _post_fix_inline_dicts(doc)
    
    # Выделяем технологии из вакансии жирным шрифтом
    if vacancy_text:
        technologies = _extract_technologies_from_vacancy(vacancy_text)
        _highlight_technologies_in_text(doc, technologies)

    # --- Новое: раздел «Примечание», если есть utochnenie ---
    _render_primichanie(doc, color, hsize, utochnenie)

    name_for_file = (cnt.get("fio") or {}).get("full_name") or "Name"
    date_str = datetime.now().strftime("%Y-%m-%d")
    # Use current directory instead of Linux path
    import os
    fn = os.path.join(os.getcwd(), f"WhiteLabel_Resume_{name_for_file.replace(' ', '_')}_{date_str}.docx")
    doc.save(fn)
    return fn

def parse_json_loose(raw):
    if isinstance(raw, dict):
        return raw
    if not isinstance(raw, str):
        raise ValueError("Ожидалась строка с JSON.")
    
    s = raw.strip()
    print(f"DEBUG: Original response length: {len(s)}")
    
    # Check for truncated JSON (incomplete braces)
    open_braces = s.count('{')
    close_braces = s.count('}')
    if open_braces > close_braces:
        print(f"WARNING: JSON appears to be truncated. Open braces: {open_braces}, Close braces: {close_braces}")
    
    # Remove markdown code blocks
    if s.startswith("```json"):
        s = s[len("```json"):].strip()
    if s.endswith("```"):
        s = s[:-3].strip()
    
    # Try parsing as-is
    try:
        return json.loads(s)
    except json.JSONDecodeError as e:
        print(f"DEBUG: First parse attempt failed: {e}")
    
    # Try extracting JSON from braces
    try:
        start, end = s.find("{"), s.rfind("}")
        if start != -1 and end != -1 and end > start:
            json_part = s[start:end+1]
            print(f"DEBUG: Extracted JSON part length: {len(json_part)}")
            return json.loads(json_part)
    except json.JSONDecodeError as e:
        print(f"DEBUG: Second parse attempt failed: {e}")
    
    # Remove invisible characters
    s2 = re.sub(r"[\u200b-\u200f\u202a-\u202e]", "", s)
    
    # Try one more time with cleaned string
    try:
        return json.loads(s2)
    except json.JSONDecodeError as e:
        print(f"DEBUG: Final parse attempt failed: {e}")
        print(f"DEBUG: Problematic JSON around error position:")
        error_pos = getattr(e, 'pos', 0)
        start_pos = max(0, error_pos - 100)
        end_pos = min(len(s2), error_pos + 100)
        print(f"DEBUG: Context: {repr(s2[start_pos:end_pos])}")
        
        # Try to fix common JSON issues
        s3 = s2
        # Fix trailing commas
        s3 = re.sub(r',(\s*[}\]])', r'\1', s3)
        # Fix unescaped quotes in strings
        s3 = re.sub(r'(?<!\\)"(?=.*".*:)', r'\\"', s3)
        
        # Try to complete truncated JSON by adding missing closing braces
        open_braces = s3.count('{')
        close_braces = s3.count('}')
        if open_braces > close_braces:
            missing_braces = open_braces - close_braces
            print(f"DEBUG: Attempting to complete truncated JSON by adding {missing_braces} closing braces")
            s3 = s3.rstrip(',\n\r\t ') + '}' * missing_braces
        
        try:
            return json.loads(s3)
        except json.JSONDecodeError as e2:
            print(f"DEBUG: Even after fixes, parsing failed: {e2}")
            # Save the problematic JSON to a file for inspection
            with open("debug_json_error.txt", "w", encoding="utf-8") as f:
                f.write(f"Original error: {e}\n")
                f.write(f"Error position: {error_pos}\n")
                f.write(f"Problematic JSON:\n{s2}")
            raise ValueError(f"Не удалось распарсить JSON. Ошибка: {e}. Детали сохранены в debug_json_error.txt")

def _extract_technologies_from_vacancy(vacancy_text: str) -> list:
    """Извлекает технологии и ключевые слова из текста вакансии"""
    import re
    
    # Общие технологии и фреймворки
    tech_patterns = [
        r'\b(?:Python|Java|JavaScript|TypeScript|C\#|C\+\+|PHP|Ruby|Go|Rust|Kotlin|Swift)\b',
        r'\b(?:React|Angular|Vue|Django|Flask|Spring|Laravel|Express|Node\.js)\b',
        r'\b(?:PostgreSQL|MySQL|MongoDB|Redis|Elasticsearch|Oracle|SQL Server)\b',
        r'\b(?:Docker|Kubernetes|AWS|Azure|GCP|Jenkins|GitLab|GitHub)\b',
        r'\b(?:Linux|Windows|MacOS|Ubuntu|CentOS)\b',
        r'\b(?:Git|SVN|Mercurial)\b',
        r'\b(?:REST|GraphQL|API|JSON|XML|SOAP)\b',
        r'\b(?:HTML|CSS|SASS|LESS|Bootstrap|Tailwind)\b',
        r'\b(?:Webpack|Vite|Babel|ESLint|Prettier)\b',
        r'\b(?:Terraform|Ansible|Puppet|Chef)\b',
        r'\b(?:Prometheus|Grafana|ELK|Splunk)\b',
        r'\b(?:Kafka|RabbitMQ|ActiveMQ)\b',
        r'\b(?:Hadoop|Spark|Airflow|Greenplum)\b'
    ]
    
    technologies = set()
    for pattern in tech_patterns:
        matches = re.findall(pattern, vacancy_text, re.IGNORECASE)
        technologies.update(matches)
    
    return list(technologies)

def _highlight_technologies_in_text(doc: Document, technologies: list):
    """Выделяет технологии жирным шрифтом в документе"""
    if not technologies:
        return
        
    import re
    
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        original_text = paragraph.text
        has_matches = False
        
        # Проверяем, есть ли технологии в этом параграфе
        for tech in technologies:
            if re.search(r'\b' + re.escape(tech) + r'\b', original_text, re.IGNORECASE):
                has_matches = True
                break
        
        if not has_matches:
            continue
            
        # Очищаем параграф и пересоздаем с выделением
        paragraph.clear()
        
        remaining_text = original_text
        while remaining_text:
            # Находим ближайшее совпадение
            earliest_match = None
            earliest_pos = len(remaining_text)
            matched_tech = None
            
            for tech in technologies:
                match = re.search(r'\b' + re.escape(tech) + r'\b', remaining_text, re.IGNORECASE)
                if match and match.start() < earliest_pos:
                    earliest_pos = match.start()
                    earliest_match = match
                    matched_tech = tech
            
            if earliest_match is None:
                # Нет больше совпадений, добавляем остальной текст
                paragraph.add_run(remaining_text)
                break
            
            # Добавляем текст до совпадения
            if earliest_pos > 0:
                paragraph.add_run(remaining_text[:earliest_pos])
            
            # Добавляем совпадение жирным
            bold_run = paragraph.add_run(earliest_match.group())
            bold_run.bold = True
            
            # Продолжаем с оставшимся текстом
            remaining_text = remaining_text[earliest_match.end():]

def build_prompt_simple(candidate_text: str, vacancy_text: str) -> str:
    return f"""
Верни ТОЛЬКО JSON {{"config":{{...}}, "content":{{...}}}} на русском, без комментариев.
White Label: не включай контакты и email. Сохрани ВСЁ содержание без сокращений.
Если нет Summary — создай 3–5 предложений. Определи должность по вакансии.
ГРЕЙД: определи только как Senior, Middle или Junior на основе опыта работы.
ГРАЖДАНСТВО: определи из локации и укажи как РФ (для России/Москвы), РБ (для Беларуси/Минска), или возьми из резюме если указано.
ПРОЕКТЫ обязательны: найди все даже если они спрятаны в обязанностях/Обо мне.
Схема:
{{
 "config": {{
   "output_format": "docx",
   "font_family": "Times New Roman",
   "font_size_main": 12,
   "font_size_headings": 14,
   "color_headings": "#1F4E79",
   "language": "ru",
   "sections": [
     "ФИО","РЕЗЮМЕ","Краткое описание профиля",
     "Ключевые навыки","Опыт работы","Образование","Дополнительная информация","Проекты"
   ],
   "white_label": true,
   "exclude_contacts": true,
   "exclude_email": true
 }},
 "content": {{
   "fio": {{"full_name":"","location":"","citizenship":"","birth_date":""}},
   "position_grade":"", "grade":"", "summary":"",
   "skills": {{}},
   "experience":[{{"company":"","position":"","period":"","responsibilities":[],"technologies":[],"achievements":[]}}],
   "education":[{{"institution":"","degree":"","years":"","details":""}}],
   "extra": [],
   "projects":[{{"title":"","role":"","period":"","description":"","technologies":[],"results":""}}]
 }}
}}

ИСХОДНОЕ РЕЗЮМЕ:
{candidate_text}

ТРЕБОВАНИЯ ВАКАНСИИ:
{vacancy_text}
""".strip()

def generate_payload_once(api_key: str,
                          candidate_text: str,
                          vacancy_text: str,
                          model_name: str = "gemini-2.5-pro") -> dict:
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    prompt = build_prompt_simple(candidate_text, vacancy_text)
    resp = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            temperature=0.1,
            max_output_tokens=12288,  # Increased from 6144 to handle larger resumes
        )
    )
    raw = _extract_text_from_gemini_response(resp)
    if not raw:
        finish = None
        safety = None
        try:
            if resp.candidates:
                finish = getattr(resp.candidates[0], "finish_reason", None)
                safety = getattr(resp.candidates[0], "safety_ratings", None)
        except Exception:
            pass
        raise ValueError(f"Пустой ответ от Gemini. finish_reason={finish}, safety={safety}")
    data = parse_json_loose(raw)
    if not isinstance(data, dict) or "config" not in data or "content" not in data:
        raise ValueError("Модель не вернула JSON с ключами {config, content}.")
    cfg = data.setdefault("config", {})
    if "sections" in cfg and "Проекты" not in cfg["sections"]:
        cfg["sections"].append("Проекты")
    return data

def create_white_label_resume_once(api_key: str,
                                   candidate_text: str,
                                   vacancy_text: str,
                                   utochnenie=None):
    payload = generate_payload_once(api_key, candidate_text, vacancy_text)
    filename = render_resume_docx(payload, vacancy_text, utochnenie=utochnenie)
    return filename

#===== Пример использования (раскомментируй, подставь API ключ и тексты) =====

import os
from dotenv import load_dotenv
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if __name__ == "__main__":
    candidate = '''
      
Contact
 Apt 321, 31, Veshnyakovskaya
 str., Moscow, 111538, Russian
 Federation.
 7-916-685-9607 (Mobile)
 sergts@mail.ru
 www.linkedin.com/in/sergts
 (LinkedIn)
 github.com/itboss2 (Portfolio)
 Top Skills
 System Architects
 Telecommunications Billing
 WebSphere ESB
 Languages
 English - C1 Certified (Full
 Professional)
 Bulgarian (Elementary)
 Russian - C2 Certifed (Native or
 Bilingual)
 Byelorussian - C2 Certified (Full
 Professional)
 Polish (Limited Working)
 German (Elementary)
 Ukrainian (Elementary)
 Certifications
 Palo Alto Networks Accredited
 Configuration Engineer (ACE) 
PAN-OS 6.1
 Cisco Network Support Engineer
 VMware Technical Sales
 Professional 5 for 6 competencies:
 Infrastructure Virtualization, Desktop
 Virtualization, Business Continuity,
 Virtualization of Business Critical
 Applications, Infrastructure as a
 Service, Management
 IBM Certified Database Associate 
DB2 10.1 Fundamentals
 Honors-Awards
 3rd prize winner of Russia Huawei
 2017 Channel Partner Skill
 Competition
 Sergey Tsuprikov, Tech
 Advisor, IBM DS, LSSBB,
 MCP,Mentor
 IT architect and project manager with 10+ years of experience in
 Big Data, Data Science, enterprise software design area | Banking
 | Telecom | FinTech | Data Lake | DWH | BI | Data Driven | Data
 Management | Data Quality
 Moscow, Moscow City, Russia
 Summary
 12+ years of experience as an IT architect for enterprise software
 (like OEBSl) design, complicated cross-platform data migration and
 integration, distributed data centers turn-key design.
 7+ years of experience as a project manager for projects up to 15M
 USD. 
7+ years of experience as a business and/or system analyst, mostly
 in the banking area. 
5+ years of experience as an out-staff instructor (PMI PMBOK,
 PRINCE2, ITIL, IBM solutions).
 10+ years of experience as a subject matter expert (SME) in banking
 and finance, sales and marketing, logistics, manufacturing, retail,
 FMCG. 
5+ years of successful daily work face-to-face and remote with
 Indian banking analysts from Oracle Corp. 
My mobile phone: 7-916-685-9607 (9:00-22:00 GMT+3, Moscow),
 Telegram: @sergts1 .
 My private email - sergts@mail.ru (please, use bossit@gmail.com
 only in case of troubles with sending to sergts@mail.ru).
 I have 30+ years of experience in IT. Have received 120+ worldwide
 IT certificates after exams (i.e. 70+ technical from IBM): Program
 & Portfolio Management Expert (#35), Project Management
 Expert (PME) #000412, IBM Data Science Professional, VMware
 Certified Professional 4 (#63533), EMCTAe, Six Sigma Black Belt
 Professional, MCP, etc. 
 Page 1 of 9
  ...
'''
    vacancy = '''
    BD-10128 (https://t.me/omega_vacancy_bot?start=3093_BD-10128)
📅 Дата публикации: 08.10.2025 12:13
... (укорочено для примера) ...
'''
    # # Пример уточнений:
    # utochnenie = [
    #     "владение французским языком (уровень не указан)",
    #     "опыт с Greenplum (подтвердить стек)",
    #     "сертификация AWS (подтвердить год)"
    # ]
    # fn = create_white_label_resume_once(GEMINI_API_KEY, candidate, vacancy, utochnenie=utochnenie)
    # print("Готово:", fn)
