from docx import Document
from PyPDF2 import PdfReader
from aiogram import Bot
import os
from gpt_gimini import sverka_vac_and_resume_json, generate_mail_for_candidate_finalist, generate_mail_for_candidate_utochnenie, generate_mail_for_candidate_otkaz, generate_cover_letter_for_client
import asyncio

from striprtf.striprtf import rtf_to_text
from dotenv import load_dotenv
import textract
import logging

from telethon_bot import ADMIN_ID
load_dotenv()

# Отключаем логирование pdfminer
logging.getLogger('pdfminer').setLevel(logging.WARNING)
logging.getLogger('pdfminer.pdfinterp').setLevel(logging.WARNING)
logging.getLogger('pdfminer.pdfpage').setLevel(logging.WARNING)
logging.getLogger('pdfminer.pdfdocument').setLevel(logging.WARNING)



CLIENT_CHANNEL = os.getenv('CLIENT_CHANNEL')

def process_doc(path: str) -> str:
    """
    Извлекает текст из .doc (старый формат Word 97–2003) с помощью textract.
    Возвращает очищенный текст без пустых строк.
    """
    try:
        text = textract.process(path).decode("utf-8", errors="ignore")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except FileNotFoundError:
        print(f"⚠️ Файл не найден: {path}")
        return ""
    except textract.exceptions.ShellError as e:
        print(f"❌ Ошибка textract при обработке {path}: {e}")
        return ""
    except Exception as e:
        print(f"⚠️ Ошибка при чтении DOC-файла {path}: {e}")
        return ""


# PDF → текст
def process_pdf(path: str) -> str:
    """
    Надёжное извлечение текста из PDF:
    1) pdfminer.six
    2) PyPDF2/pypdf (с попыткой strict=False, если доступно)
    3) Ремонт PDF через pikepdf и повторная попытка (pdfminer/textract)
    Возвращает очищенный текст без пустых строк.
    """
    def _clean(txt: str) -> str:
        return "\n".join([ln.strip() for ln in (txt or "").splitlines() if ln.strip()]).strip()

    # --- 1) pdfminer.six ---
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract_text
        txt = _clean(pdfminer_extract_text(path) or "")
        # если текста достаточно — этого хватает
        if len(txt) > 200:
            return txt
        # иначе не выходим — попробуем другие способы (может быть скан или «кривой» PDF)
    except Exception as e:
        print(f"⚠️ pdfminer.six не справился: {e}")

    # --- 2) PyPDF2 / pypdf ---
    try:
        from PyPDF2 import PdfReader
        try:
            # pypdf 3.x: параметр strict отсутствует
            reader = PdfReader(path)
        except TypeError:
            # PyPDF2 1.x/2.x: можно ослабить строгость
            reader = PdfReader(path, strict=False)
        pages_text = []
        for p in reader.pages:
            t = p.extract_text() or ""
            if t.strip():
                pages_text.append(t)
        txt = _clean("\n".join(pages_text))
        if txt:
            return txt
    except Exception as e:
        # именно ваш кейс
        if "Odd-length string" in str(e):
            print("⚠️ PyPDF2: Odd-length string — попробую отремонтировать PDF через pikepdf…")
        else:
            print(f"⚠️ PyPDF2/pypdf упал: {e}")

    # --- 3) Ремонт через pikepdf и повтор ---
    try:
        import tempfile, pikepdf
        with pikepdf.open(path) as pdf:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                pdf.save(tmp.name)
                repaired_path = tmp.name

        # снова попробуем pdfminer
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract_text
            txt = _clean(pdfminer_extract_text(repaired_path) or "")
            if txt:
                return txt
        except Exception as e:
            print(f"⚠️ pdfminer после ремонта не справился: {e}")

        # финальный фоллбэк: textract (может дернуть tesseract, если установлен)
        try:
            import textract
            raw = textract.process(repaired_path).decode("utf-8", errors="ignore")
            txt = _clean(raw)
            return txt
        except Exception as e:
            print(f"❌ textract тоже не смог: {e}")

    except Exception as e:
        print(f"❌ Ремонт PDF через pikepdf не удался: {e}")

    return ""


# DOCX → текст
def process_docx(path: str) -> str:
    """
    Извлекает весь текст из .docx, включая таблицы и вложенные ячейки.
    Возвращает объединённый текст.
    """
    try:
        doc = Document(path)
        texts = []

        # --- Параграфы ---
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text.strip())

        # --- Таблицы ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)

        # Удаляем дубликаты и объединяем
        text = "\n".join(dict.fromkeys(texts))
        return text.strip()

    except Exception as e:
        print(f"❌ Ошибка чтения DOCX: {e}")
        return ""

# RTF → текст
def process_rtf(path: str) -> str:
    """
    Читает RTF-файл и возвращает чистый текст.
    Работает без Pandoc.
    """
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    text = rtf_to_text(content)
    return text

# TXT → текст
def process_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

async def process_file_and_gpt(path: str, bot: Bot, user_id: int|str, vac_text: str, file_name: str):
    ext = path.split(".")[-1].lower()
    
    try:
        if ext == "pdf":
            text = process_pdf(path)
        elif ext == "docx":
            text = process_docx(path)
        elif ext == "doc":
            text = process_doc(path)
        elif ext == "rtf":
            text = process_rtf(path)
        elif ext == "txt":
            text = process_txt(path)
        else:
            await bot.send_message(user_id, f"⚠️ Формат {ext} не поддерживается: {path}")
            return
        
        data  = await background_sverka(resume_text=text, vacancy_text=vac_text, bot=bot, user_id=user_id, file_name=file_name)
        
        os.remove(path)
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка в {path}: {e}")
    finally:
        return data or None
        
async def background_sverka(resume_text: str, vacancy_text: str, bot: Bot, user_id: int|str, file_name: str):
    try:
        result_gpt = await sverka_vac_and_resume_json(resume_text, vacancy_text, file_name)
        
        if result_gpt:
            result = display_analysis(result_gpt)
            result_gpt = clean_json(result_gpt)
            verdict = result_gpt.get("summary").get("verdict")
            candidate = result_gpt.get("candidate").get("full_name")
            
            
            return {'candidate': candidate, 'verdict': verdict, 'sverka_text': result, 'candidate_json': result_gpt, 'resume_text': resume_text}
        else:
            await bot.send_message(ADMIN_ID, "❌ Ошибка при сверке вакансии")
    except Exception as e:
        await bot.send_message(ADMIN_ID, f"🔥 Ошибка при сверке: {e}")
        return None
    
    
        
        
        

import json


def clean_json(json_data):
    if isinstance(json_data, str):
        clean_str = json_data.strip()
        if clean_str.startswith('```json'):
            clean_str = clean_str[len('```json'):].strip()
        if clean_str.endswith('```'):
            clean_str = clean_str[:-len('```')].strip()
        
        try:
            data = json.loads(clean_str)
        except json.JSONDecodeError:
            return "Ошибка: Некорректный формат JSON после очистки."
    else:
        data = json_data
    return data

def display_analysis(json_data):
    """
    Принимает JSON-строку или словарь Python и ВОЗВРАЩАЕТ
    структурированный отчет, содержащий Имя кандидата, "Таблицу соответствия" и "Итог".
    Если поле отсутствует, выводит '❌'.
    Автоматически удаляет маркеры блока кода ```json и ```.
    """
    processed_data = json_data
    output_lines = []  # Список для хранения всех строк отчета

    # --- Блок очистки входных данных ---
    processed_data = clean_json(processed_data)
    data = processed_data

    # Вспомогательная функция для форматирования поля "ключ: значение"
    def format_field(key, value):
        val_str = value if value else "❌"
        return f"{key}: {val_str}"
    location = data.get("candidate", {}).get('location')
    city = location.get('city', None)
    country = location.get('country', None)
    if city == 'Нет (требуется уточнение)':
        city = None
    if country == 'Нет (требуется уточнение)':
        country = None
    if city and country:
        location = f"{city}, {country}"
    elif city:
        location = city
    elif country:
        location = country
    else:
        location = "не указано"
    # --- КАНДИДАТ (только ФИО) ---
    output_lines.append("="*15 + " 👤 КАНДИДАТ " + "="*15)
    candidate = data.get("candidate", {})
    output_lines.append(format_field("ФИО", candidate.get('full_name')))
    output_lines.append(format_field("—Дата рождения", candidate.get('birth_date').get('date')))
    output_lines.append(format_field("—Зарплатные ожидания", data.get('summary').get('salary_expectations')))
    output_lines.append(format_field("—Локация", location))
    output_lines.append(format_field("—Стек технологий", ", ".join(candidate.get('tech_stack'))) )


    # --- ТАБЛИЦА СООТВЕТСТВИЯ ---
    output_lines.append("\n" + "="*12 + " ✅ ТАБЛИЦА СООТВЕТСТВИЯ " + "="*12)
    compliance = data.get("compliance_check", {})
    status_map = { "Да": "✅", "Нет (требуется уточнение)": "⚠️", "Нет (точно нет)": "❌" }
    
    must_haves = compliance.get('must_have')
    if must_haves:
        for req in must_haves:
            icon = status_map.get(req.get('status'), '▫️')
            if req.get('status') == "Нет (требуется уточнение)" or req.get('status') == "Нет (точно нет)":
                output_lines.append(f"    {icon} {req.get('requirement')}")
                output_lines.append(f"({req.get('comment').replace('⚠️', '').replace('❌', '')})\n")
            else:
                output_lines.append(f"    {icon} {req.get('requirement')}\n")


    nice_to_haves = compliance.get('nice_to_have')
    if nice_to_haves:
        for req in nice_to_haves:
            icon = status_map.get(req.get('status'), '▫️')
            if req.get('status') == "Нет (требуется уточнение)" or req.get('status') == "Нет (точно нет)":
                output_lines.append(f"    {icon} {req.get('requirement')}")
                output_lines.append(f"({req.get('comment').replace('⚠️', '').replace('❌', '')})\n")
            else:
                output_lines.append(f"    {icon} {req.get('requirement')}\n")   

    # --- ИТОГ ---
    output_lines.append("\n" + "="*17 + " 🏁 ИТОГ " + "="*17)
    summary = data.get("summary", {})
    if summary:
        output_lines.append(format_field("Вердикт", summary.get('verdict')))
    output_lines.append("="*41)

    return "\n".join(output_lines)




def create_finalists_table(finalists: list[dict]):
  """
  Создает таблицу финалистов в формате Markdown.

  Args:
    finalists: Список словарей, где каждый словарь представляет финалиста
               с ключами 'name', 'grade', 'location', 'stack', и 'salary'.

  Returns:
    Строка с таблицей в формате Markdown.
  """

  
  
  header = "| ФИО/ФИ | Грейд | Локация | Ключевой стек | Зарплатные ожидания |\n"
  separator = "|---|---|---|---|---|\n"
  body = ""
  for finalist in finalists:
    if isinstance(finalist, str):
      continue
    candidate = finalist.get("candidate", {})
    summary = finalist.get("summary", {})
    verdict = summary.get("verdict", "")
    if verdict == "Полностью подходит":
      body += f"| {candidate['full_name'] or '❌'} | {candidate['grade_and_position'] or '❌'} | {candidate['location']['city'] or '❌'} | {summary['salary_expectations'] or '❌'} |{summary['verdict'] or '❌'}\n"
    elif verdict == "Частично подходит (нужны уточнения)":
      body += f"| {candidate['full_name'] or '❌'} | {candidate['grade_and_position'] or '❌'} | {candidate['location']['city'] or '❌'} | {summary['salary_expectations'] or '❌'} |{summary['verdict'] or '❌'}\n"
    elif verdict == "Не подходит":
      body += f"| {candidate['full_name'] or '❌'} | {candidate['grade_and_position'] or '❌'} | {candidate['location']['city'] or '❌'} | {summary['salary_expectations'] or '❌'} |{summary['verdict'] or '❌'}\n"
  return header + separator + body



    
    
async def create_mails(finalist: dict, user_name: str,vacancy: str, group_id: int, thread_id: int, verdict: str):
    try:
    
      if isinstance(finalist, str):
        print("❌ Неверный формат данных финалиста")
        return None
      
      print(verdict)
      if verdict == "PP":
        res = await generate_mail_for_candidate_finalist(finalist, user_name, group_id, thread_id)
        return res
      elif verdict == "CP":
        res = await generate_mail_for_candidate_utochnenie(finalist, user_name, vacancy, group_id, thread_id)
        return res
      elif verdict == "NP":
        res = await generate_mail_for_candidate_otkaz(finalist, user_name)
        return res
    except Exception as e:
      print(f"❌ Произошла ошибка при создании письма: {e}")
      return None